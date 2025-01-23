import os
import subprocess
from docx import Document
from docxcompose.composer import Composer
class DocumentProcessor:
    def __init__(self):
        """
        Initialize the DocumentProcessor class.
        """
        pass

    def extract_latex_content(self, input_file, output_file):
        """
        Extract LaTeX content between ```latex and ``` markers from an input file
        and save it to a new output file.
        
        Args:
            input_file (str): Path to the input file
            output_file (str): Path to the output file
        
        Returns:
            bool: True if extraction successful, False otherwise
        """
        try:
            with open(input_file, 'r', encoding='utf-8') as infile:
                lines = infile.readlines()
            
            # Find the start and end of LaTeX content
            start_index = -1
            end_index = -1
            
            for i, line in enumerate(lines):
                if line.strip() == '```latex':
                    start_index = i + 1
                elif start_index != -1 and line.strip() == '```':
                    end_index = i
                    break
            
            # Check if LaTeX content was found
            if start_index == -1 or end_index == -1:
                print("No LaTeX content found between ```latex and ``` markers.")
                return False
            
            # Extract and save LaTeX content
            latex_content = lines[start_index:end_index]
            
            with open(output_file, 'w', encoding='utf-8') as outfile:
                outfile.writelines(latex_content)
            
            print(f"LaTeX content successfully extracted to {output_file}")
            return True
        
        except Exception as e:
            print(f"An error occurred during LaTeX extraction: {e}")
            return False

    def conversion_to_office(self, file, output_format):
        """
        Convert a file to Office format using pandoc.
        
        Args:
            file (str): Input file path
            output_format (str): Desired output format (e.g., 'docx', 'pdf')
        
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            # Construct the output filename
            output_file = f"{os.path.splitext(file)[0]}.{output_format}"
            
            # Construct and run the pandoc command
            command = f"pandoc '{file}' -o '{output_file}'"
            result = subprocess.run(command, shell=True, check=True, 
                                    capture_output=True, text=True)
            
            print(f"Successfully converted {file} to {output_file}")
            return True
        
        except subprocess.CalledProcessError as e:
            print(f"Conversion error: {e}")
            print(f"Standard Output: {e.stdout}")
            print(f"Standard Error: {e.stderr}")
            return False
        except Exception as e:
            print(f"An unexpected error occurred during conversion: {e}")
            return False

    def insert_content_into_template(self, content_docx, template_docx, output_docx):
        """
        Insert content from one DOCX file into a template DOCX file.
        
        Args:
            content_docx (str): Path to the source DOCX file
            template_docx (str): Path to the template DOCX file
            output_docx (str): Path for the output DOCX file
        
        Returns:
            bool: True if insertion successful, False otherwise
        """
        try:
            # Open documents
            template_docx = Document(template_docx)
            content_docx = Document(content_docx)
            #Creation of a composer instance
            composer = Composer(template_docx)
            #Break page after template
            #template_docx.add_page_break()         
            # Append source document to template
            composer.append(content_docx)
            #for element in content_docx.element.body:
            #	template_docx.element.body.append(element)
            composer.save(output_docx)
            #template_docx.save(output_docx)          
            print(f"Content successfully inserted into {output_docx}")
            return True
        
        except Exception as e:
            print(f"An error occurred during document insertion: {e}")
            return False
    
    def insert_string_into_docx(self,template_path, string_to_insert, output_path,align=None):
        """
        Insert a string into a template DOCX file.
    
        Args:
        template_path (str): Path to the template DOCX file
        string_to_insert (str): String to insert into the document
        output_path (str): Path for the output DOCX file
        """
        # Open the template document
        doc = Document(template_path)
    
        # Add a page break
        #doc.add_page_break()
    
        # Add the string as a paragraph
        p=doc.add_paragraph(string_to_insert)
        if align!=None:
        	p.alignment = align # for left, 1 for center, 2 right, 3 justify ....doc.add_paragraph(string_to_insert,jc=align)
#        paragraph.alignment = 0 # for left, 1 for center, 2 right, 3 justify ....
    
        # Save the modified document
        doc.save(output_path)
